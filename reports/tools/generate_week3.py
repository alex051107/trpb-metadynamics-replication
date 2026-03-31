"""Generate Week 3 weekly report as .docx"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Global style --
style = doc.styles["Normal"]
font = style.font
font.name = "Arial"
font.size = Pt(11)
font.color.rgb = RGBColor(0, 0, 0)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = "Arial"
        if level == 1:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
    return h

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(10)
    # Rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
    return table

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.bold = bold
    return p

# ============================================================
# HEADER
# ============================================================
add_table(
    ["Field", "Value"],
    [
        ["To", "Prof. Anima Anandkumar (aNima Lab, Caltech)"],
        ["From", "Zhenpeng Liu"],
        ["Date", "2026-03-28"],
        ["Week", "#3: HPC environment setup and MetaDynamics toolchain validation"],
    ],
)
doc.add_paragraph()

# ============================================================
# SUMMARY
# ============================================================
add_heading("Summary", level=1)
add_para(
    "This week I shifted from paper reading to hands-on work. "
    "I set up the full simulation stack on UNC Longleaf (AMBER 24p3, GROMACS 2026.0, PLUMED 2.9) "
    "and ran an alanine dipeptide well-tempered MetaDynamics test to validate the toolchain. "
    "The test produced a physically reasonable free energy surface with expected Ramachandran basins. "
    "The main unresolved item is PLP cofactor parameterization: the JACS 2019 SI does not include "
    "the actual force field files, so I will need to generate them from scratch using antechamber and Gaussian."
)

# ============================================================
# SECTION 1: WORK DONE
# ============================================================
add_heading("Section 1: Work Done", level=1)
add_table(
    ["#", "Task", "Category", "Context", "Status"],
    [
        ["1", "Read WIREs 2020 (Osuna review)", "Paper", "DOI: 10.1002/wcms.1500", "Done"],
        ["2", "Read Protein Science 2025 (standalone TrpB)", "Paper", "DOI: 10.1002/pro.70103", "Done"],
        ["3", "Install PLUMED 2.9 on Longleaf", "Setup", "conda-forge, trpb-md env", "Done"],
        ["4", "Install GROMACS 2026.0 with PLUMED patch", "Setup", "conda-forge, PLUMED support confirmed", "Done"],
        ["5", "Verify AMBER 24p3 (pmemd.cuda)", "Setup", "module load amber/24p3", "Done"],
        ["6", "Extract MetaDynamics params from JACS 2019 SI", "Analysis", "ja9b03646_si_001.pdf", "Done"],
        ["7", "Fact-check all project documentation", "Analysis", "5 parallel checks, 7 errors corrected", "Done"],
        ["8", "Download PDB structures (1WDW, 3CEP, 4HPX, 5IXJ)", "Setup", "Path CV endpoints + references", "Done"],
        ["9", "Run alanine dipeptide WT-MetaD toy test", "Simulation", "6.1 ns, 6088 hills, Job 39960327", "Done"],
        ["10", "Generate and analyze FES from toy test", "Analysis", "plumed sum_hills, matplotlib", "Done"],
        ["11", "Verify Gaussian16/ORCA availability", "Setup", "PLP RESP parameterization feasible", "Done"],
        ["12", "Set up 6-stage replication pipeline", "Documentation", "Profiler to Journalist", "Done"],
    ],
)

# ============================================================
# SECTION 2: WHAT I LEARNED
# ============================================================
add_heading("Section 2: What I Learned", level=1)

add_heading("Protein Science 2025: a natural standalone TrpB exists", level=2)
add_para(
    "Kinateder et al. found a naturally occurring standalone TrpB from Pelodictyon luteolum (plTrpB) "
    "that contains the same six key residues (LBCA-Res6) identified through ancestral sequence reconstruction. "
    "Its kcat is 0.35 s^-1 without TrpA, compared to 0.009 s^-1 for a variant where those six residues "
    "are replaced by consensus residues (Kinateder et al., Protein Science 2025, Table 1)."
)
add_para(
    "The COMM domain O-to-C barrier in plTrpB is 5 to 6 kcal/mol. "
    "This is higher than the 2 kcal/mol barrier in PfTrpB-0B2 (Maria-Solano et al., JACS 2019, Figure 3) "
    "but still low enough for standalone activity. "
    "This gives me a natural benchmark: if GenSLM-230 truly has higher activity than 0B2, "
    "its barrier should be below 2 kcal/mol."
)

add_heading("GROMACS + PLUMED2 toolchain: validated on alanine dipeptide", level=2)
add_para(
    "I ran a well-tempered MetaDynamics simulation of alanine dipeptide at 300 K using "
    "GROMACS 2026.0 + PLUMED 2.9 on Longleaf. "
    "The simulation deposited 6088 Gaussian hills over 6.1 ns. "
    "The Gaussian height decayed from ~0.6 kJ/mol at the start to 0.19 kJ/mol at the end, "
    "confirming that the well-tempered scheme is working correctly."
)
add_para(
    "The resulting FES (Figure 1) shows the expected Ramachandran basins. "
    "The alpha_R basin appears at phi = -53, psi = -32 with a free energy of 6.1 kJ/mol. "
    "The alpha_L basin is at phi = 74, psi = 18 at 3.3 kJ/mol. "
    "The global minimum is in the positive-phi region (phi = 67, psi = -152), "
    "which is a standard result for alanine dipeptide at 300 K."
)

add_heading("Parameter extraction from SI: what the paper gives and what it does not", level=2)
add_para(
    "I extracted all MetaDynamics parameters from the JACS 2019 Supporting Information: "
    "Gaussian height = 0.15 kcal/mol (SI Methods), deposition pace = 2 ps, bias factor = 10, "
    "adaptive Gaussian width, 10 walkers per system, 50 to 100 ns per walker (SI Table S2). "
    "The force field is ff14SB with TIP3P water and 8 A PME cutoff."
)
add_para(
    "One critical gap: the SI does not include the actual PLP cofactor parameter files (mol2/frcmod). "
    "It only states that GAFF + RESP charges at HF/6-31G(d) were used via antechamber. "
    "I confirmed that Gaussian 16c02 and ORCA 6.0/6.1 are available on Longleaf, "
    "so I can regenerate these parameters from scratch."
)

# Insert FES figure
fes_path = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "replication", "analysis", "toy-alanine", "fes_alanine.png"
)
if os.path.exists(fes_path):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(fes_path, width=Inches(4.5))
    cap = doc.add_paragraph("Figure 1. Free energy surface of alanine dipeptide from well-tempered MetaDynamics "
                            "(6.1 ns, 300 K, ff14SB, GROMACS 2026.0 + PLUMED 2.9). "
                            "Energy capped at 40 kJ/mol for visualization.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.name = "Arial"
        run.italic = True

# ============================================================
# SECTION 3: PROBLEMS ENCOUNTERED
# ============================================================
add_heading("Section 3: Problems Encountered and How I Solved Them", level=1)

add_heading("PDB atom names not matching force field definitions", level=2)
add_para(
    "What happened: pdb2gmx rejected the alanine dipeptide PDB with "
    '"Atom N in residue ACE not found in rtp entry."'
)
add_para(
    "Root cause: the ACE capping group in AMBER ff14SB only contains CH3, C, and O atoms. "
    "The initial PDB incorrectly included an N atom in ACE."
)
add_para(
    "How I fixed it: I rewrote the PDB with the correct atom set for ACE (CH3, C, O) "
    "and NME (N, H, CH3), matching the aminoacids.rtp definitions."
)
add_para(
    "Lesson: never write PDB atom names from memory. "
    "Always check the force field rtp file or run a pdb2gmx dry run first."
)

add_heading("PLUMED atom indices wrong after pdb2gmx reordering", level=2)
add_para(
    "What happened: the PLUMED input file defined phi/psi torsions using atom indices 4, 5, 6, 7. "
    "PLUMED ran without errors, but the CV values were nonsensical."
)
add_para(
    "Root cause: pdb2gmx reorders atoms (adds hydrogens, renames). "
    "The correct indices for phi/psi after reordering were 5, 7, 9, 15 and 7, 9, 15, 17, "
    "confirmed by inspecting the .gro output file."
)
add_para(
    "Lesson: PLUMED atom indices must always be determined from the post-pdb2gmx .gro file, "
    "never from the input PDB."
)

add_heading("Slurm OMP_NUM_THREADS conflict", level=2)
add_para(
    "What happened: the first two Slurm submissions (Jobs 39960167, 39960248) failed within seconds. "
    "GROMACS reported a fatal error about inconsistent thread settings."
)
add_para(
    "Root cause: Longleaf sets OMP_NUM_THREADS=1 by default, "
    "but the GROMACS command used -ntomp 4. The mismatch caused an immediate crash."
)
add_para(
    "How I fixed it: I added 'export OMP_NUM_THREADS=4' to the Slurm script, "
    "placed after conda activation and before gmx mdrun. Job 39960327 ran successfully."
)
add_para(
    "Lesson: all future Slurm scripts must explicitly set OMP_NUM_THREADS "
    "to match the -ntomp flag."
)

# ============================================================
# SECTION 4: OPEN QUESTIONS
# ============================================================
add_heading("Section 4: Open Questions", level=1)

add_para(
    "1. PLP parameter files: the JACS 2019 SI does not provide mol2/frcmod files "
    "for the PLP cofactor intermediates. I plan to regenerate them using antechamber + Gaussian16. "
    "If this proves difficult, should I contact the Osuna lab to request their files?",
)
add_para(
    "2. D-Trp facial selectivity CV: I still have no literature precedent for defining "
    "a collective variable that captures the indole attack direction on aminoacrylate. "
    "This blocks Direction B (D-Trp screening). "
    "Is there someone in the lab who has worked on stereoselectivity CVs?",
)
add_para(
    "3. GenSLM-230 and NdTrpB structures: I need the actual sequences and/or homology models "
    "to begin Phase 2. Are these available in the lab, or should I build homology models myself?",
)
add_para(
    "4. GROMACS GPU support in conda: the conda-forge GROMACS build supports PLUMED, "
    "but I have not verified GPU acceleration ('gmx mdrun -nb gpu'). "
    "For the 10-walker MetaD runs this will matter.",
)

# ============================================================
# SECTION 5: NEXT WEEK PLAN
# ============================================================
add_heading("Section 5: Next Week Plan", level=1)

add_heading("Papers to read", level=2)
add_para(
    "JACS 2019 (Maria-Solano et al.): full read with annotation. "
    "I need to understand Figure 2 (FEL comparison) and Figure 3 (productive vs unproductive closure) "
    "in detail before I can evaluate my own FES results."
)
add_para(
    "Nature Communications 2026 (Lambert et al.): full read. "
    "Focus on how GenSLM-230 was selected and what experimental assays were used."
)

add_heading("Tasks", level=2)
add_para("1. PLP parameterization (Ain intermediate): extract PLP coordinates from 1WDW, "
         "run antechamber with GAFF + RESP at HF/6-31G(d) using Gaussian16. "
         "Depends on: Gaussian availability (confirmed).")
add_para("2. Generate 15-frame O-to-C reference path: linear interpolation of C-alpha atoms "
         "(residues 97-184 + 282-305) between 1WDW (open) and 3CEP (closed). "
         "Depends on: PDB files (downloaded).")
add_para("3. Test GROMACS GPU acceleration: run 'gmx mdrun -nb gpu' with a small system "
         "to confirm the conda build supports GPU offloading.")
add_para("4. Begin PfTrpS(Ain) system preparation: solvation, minimization, and heating protocol. "
         "Depends on: PLP parameters (task 1).")

# Save
outpath = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "WeeklyReport_Week3_2026-03-28.docx")
doc.save(outpath)
print(f"Saved: {outpath}")
print(f"Size: {os.path.getsize(outpath)} bytes")
