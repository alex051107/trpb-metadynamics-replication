"""Generate Week 2 weekly report as .docx (v2: advisor-perspective rewrite)"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# -- Global style --
style = doc.styles["Normal"]
font = style.font
font.name = "Arial"
font.size = Pt(11)
font.color.rgb = RGBColor(0, 0, 0)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = "Arial"
        run.font.size = Pt(14 if level == 1 else 12)
    return h


def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(10)
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = t.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
    return t


def para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


# ============================================================
# HEADER
# ============================================================
table(
    ["Field", "Value"],
    [
        ["To", "Prof. Anima Anandkumar and Anima Lab members"],
        ["From", "Zhenpeng Liu (UNC Chapel Hill)"],
        ["Date", "2026-03-28"],
        ["Week", "#2: Simulation environment setup and MetaDynamics toolchain validation"],
    ],
)
doc.add_paragraph()

# ============================================================
# PROJECT CONTEXT
# ============================================================
heading("Project Context", level=1)
para(
    "I am replicating a MetaDynamics-based conformational analysis of TrpB enzymes, "
    "following the protocol published by the Osuna group (JACS 2019). "
    "TrpB is an enzyme that synthesizes tryptophan. Its catalytic activity depends on "
    "whether a structural region called the COMM domain "
    "(a ~88-residue \"lid\" that opens and closes) can reach the correct closed conformation. "
    "The Osuna group showed that MetaDynamics "
    "(an enhanced sampling method that accelerates rare conformational transitions) "
    "can predict TrpB activity by mapping the free energy landscape of this open-to-closed transition."
)
para(
    "My goal is to first validate this method on known systems, then apply it to screen "
    "AI-generated enzyme variants from the Anima Lab GenSLM pipeline. "
    "This week I focused on setting up the simulation environment on UNC Longleaf HPC "
    "and running a validation test to confirm the software toolchain works."
)

# ============================================================
# PIPELINE OVERVIEW
# ============================================================
heading("Pipeline Overview", level=1)
para("Phase 1: Replicate Osuna JACS 2019 results", bold=True)
pipeline_items = [
    ("[1] Environment setup (AMBER, GROMACS, PLUMED)", "DONE (this week)"),
    ("[2] Toolchain validation (alanine dipeptide toy test)", "DONE (this week)"),
    ("[3] PLP cofactor parameterization", "NEXT"),
    ("[4] O-to-C reference path generation (15 frames)", ""),
    ("[5] PfTrpS system preparation", ""),
    ("[6] 500 ns conventional MD", ""),
    ("[7] Well-tempered MetaDynamics (10 walkers)", ""),
    ("[8] Compare FES with published Figure 2a", ""),
]
for item, status in pipeline_items:
    if status:
        para(f"  {item}    {status}")
    else:
        para(f"  {item}")
para("Phase 2: Apply to GenSLM-designed enzymes (after Phase 1 benchmark passes)", bold=True)
para("Phase 3: Build reward function for generative design loop", bold=True)

# ============================================================
# SUMMARY
# ============================================================
heading("Summary", level=1)
para(
    "This week I moved from paper reading into hands-on setup. "
    "I installed three simulation packages on UNC Longleaf: "
    "AMBER 24p3 (for conventional molecular dynamics), "
    "GROMACS 2026.0 (the dynamics engine used in the original paper), "
    "and PLUMED 2.9 (the MetaDynamics plugin). "
    "I then ran an alanine dipeptide MetaDynamics test to validate the toolchain. "
    "The test produced a free energy surface (FES: a map showing which conformations "
    "are energetically preferred) with physically reasonable energy basins. "
    "The main unresolved item is PLP cofactor parameterization: "
    "the original paper's Supporting Information does not include the actual force field files, "
    "so I will need to regenerate them using antechamber "
    "(a tool for deriving force field parameters from quantum chemistry calculations) "
    "and Gaussian (quantum chemistry software)."
)

# ============================================================
# SECTION 1: WORK DONE
# ============================================================
heading("Section 1: Work Done", level=1)
table(
    ["#", "Task", "Category", "Context", "Status"],
    [
        ["1", "Read WIREs 2020 (Osuna review)", "Paper", "DOI: 10.1002/wcms.1500", "Done"],
        ["2", "Read Protein Science 2025 (standalone TrpB)", "Paper", "DOI: 10.1002/pro.70103", "Done"],
        ["3", "Install PLUMED 2.9 on Longleaf", "Setup", "conda-forge, trpb-md env", "Done"],
        ["4", "Install GROMACS 2026.0 with PLUMED patch", "Setup", "conda-forge", "Done"],
        ["5", "Verify AMBER 24p3 (GPU)", "Setup", "module load amber/24p3", "Done"],
        ["6", "Extract MetaDynamics params from SI (Supporting Information)", "Analysis", "JACS 2019 SI PDF", "Done"],
        ["7", "Download PDB structures", "Setup", "1WDW, 3CEP, 4HPX, 5IXJ", "Done"],
        ["8", "Alanine dipeptide WT-MetaD test", "Simulation", "6.1 ns, 6088 hills", "Done"],
        ["9", "Generate and analyze FES", "Analysis", "plumed sum_hills", "Done"],
        ["10", "Verify Gaussian16 / ORCA availability", "Setup", "For PLP RESP charges", "Done"],
        ["11", "Fact-check project documentation", "Analysis", "7 errors corrected", "Done"],
    ],
)

# ============================================================
# SECTION 2: WHAT I LEARNED
# ============================================================
heading("Section 2: What I Learned", level=1)

heading("A natural standalone TrpB provides a conformational benchmark", level=2)
para(
    "To understand what a \"naturally evolved\" solution to the standalone TrpB problem "
    "looks like, I read Kinateder et al. (Protein Science 2025). "
    "They found a TrpB from Pelodictyon luteolum (plTrpB) that works without its usual partner protein TrpA. "
    "Its catalytic rate is 0.35 s^-1 on its own, compared to 0.009 s^-1 for a variant "
    "where six key residues are replaced (Table 1 in the paper)."
)
para(
    "The COMM domain open-to-closed energy barrier in plTrpB is about 5 to 6 kcal/mol. "
    "For comparison, the engineered variant PfTrpB-0B2 has a barrier of only ~2 kcal/mol "
    "(Maria-Solano et al., JACS 2019, Figure 3). "
    "This gives me a quantitative scale: if GenSLM-230 truly outperforms 0B2, "
    "I should see its barrier fall below 2 kcal/mol in my MetaDynamics simulations."
)

heading("Validating the GROMACS + PLUMED2 toolchain on alanine dipeptide", level=2)
para(
    "Before running the actual TrpB simulations, I needed to confirm that "
    "GROMACS (the molecular dynamics engine) and PLUMED "
    "(a plugin that adds MetaDynamics capabilities) "
    "work together correctly on Longleaf. "
    "I ran a well-tempered MetaDynamics simulation "
    "(a variant where the bias gradually decreases for better convergence) "
    "on alanine dipeptide, a standard test system."
)
para(
    "The simulation deposited 6,088 Gaussian hills over 6.1 ns at 300 K. "
    "The hill height decayed from ~0.6 kJ/mol at the start to 0.19 kJ/mol at the end, "
    "confirming that the well-tempered scheme works. "
    "The resulting FES (Figure 1) shows the expected energy basins on the Ramachandran map "
    "(a plot of backbone dihedral angles phi and psi). "
    "The right-handed alpha-helix basin sits at phi = -53, psi = -32 degrees "
    "(6.1 kJ/mol above the global minimum). "
    "The positive-phi region is more populated than at 300 K. "
    "This is a standard result for alanine dipeptide at 300 K."
)

heading("What the original paper provides and what it does not", level=2)
para(
    "I extracted all MetaDynamics parameters from the JACS 2019 Supporting Information. "
    "The key ones are listed in the Parameters section below. "
    "One critical gap: the SI states that PLP cofactor charges were computed using "
    "GAFF (a general-purpose small-molecule force field) with RESP charges "
    "(electrostatic charges fit to quantum calculations at the HF/6-31G(d) level), "
    "but the actual parameter files are not included. "
    "I confirmed that Gaussian 16c02 is available on Longleaf, "
    "so I can regenerate these parameters from scratch next week."
)

# Insert FES figure
base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
fes_path = os.path.join(base, "replication", "analysis", "toy-alanine", "fes_alanine.png")
if os.path.exists(fes_path):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(fes_path, width=Inches(4.5))
    cap = doc.add_paragraph(
        "Figure 1. Free energy surface of alanine dipeptide from well-tempered MetaDynamics "
        "(6.1 ns, 300 K, ff14SB, GROMACS 2026.0 + PLUMED 2.9). "
        "Energy capped at 40 kJ/mol for visualization. "
        "Stars/circles mark key energy basins."
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.name = "Arial"
        run.italic = True

# ============================================================
# SECTION 3: PROBLEMS ENCOUNTERED
# ============================================================
heading("Section 3: Problems Encountered and How I Solved Them", level=1)

heading("Force field atom name mismatch in PDB input", level=2)
para(
    "What happened: the GROMACS topology builder (pdb2gmx) rejected the alanine dipeptide PDB "
    "with the error \"Atom N in residue ACE not found in rtp entry.\""
)
para(
    "Root cause: the ACE capping group in the AMBER ff14SB force field only contains "
    "CH3, C, and O atoms. The initial PDB file incorrectly included an N atom."
)
para(
    "How I fixed it: I rewrote the PDB with the correct atom set, "
    "matching the aminoacids.rtp definitions in the force field directory."
)
para(
    "Lesson: PDB atom names must be verified against the force field definition files "
    "before running any topology builder. A quick pdb2gmx dry run catches these errors early."
)

heading("PLUMED atom indices shifted after topology generation", level=2)
para(
    "What happened: the PLUMED input file defined phi/psi torsion angles "
    "using atom indices from the original PDB. "
    "The simulation ran without errors, but the collective variable "
    "(CV: a reduced coordinate describing the molecular conformation) values were wrong."
)
para(
    "Root cause: pdb2gmx reorders atoms when adding hydrogens and renaming residues. "
    "The correct indices (5, 7, 9, 15 for phi; 7, 9, 15, 17 for psi) "
    "had to be read from the output .gro file, not the input PDB."
)
para(
    "Lesson: PLUMED atom indices must always come from the post-topology .gro file. "
    "I will build this check into all future PLUMED input generation scripts."
)

heading("Slurm thread settings conflicting with GROMACS", level=2)
para(
    "What happened: the first two Slurm job submissions failed within seconds. "
    "GROMACS reported a fatal error about inconsistent OpenMP thread counts."
)
para(
    "Root cause: Longleaf sets OMP_NUM_THREADS=1 by default. "
    "The GROMACS command used -ntomp 4. The mismatch caused an immediate crash."
)
para(
    "How I fixed it: I added \"export OMP_NUM_THREADS=4\" to the Slurm script, "
    "placed after environment activation and before the GROMACS command. "
    "The third submission ran successfully for the full walltime."
)
para(
    "Lesson: all Slurm submission scripts must explicitly set OMP_NUM_THREADS "
    "to match the requested thread count. This is now in my script template."
)

# ============================================================
# SECTION 4: OPEN QUESTIONS
# ============================================================
heading("Section 4: Open Questions", level=1)

para(
    "1. PLP parameter files: the JACS 2019 SI does not provide the actual mol2/frcmod files "
    "for the PLP cofactor. I plan to regenerate them with antechamber + Gaussian16. "
    "If this proves difficult, should I contact the Osuna lab directly?"
)
para(
    "2. D-Trp facial selectivity: I have no literature precedent for defining a collective variable "
    "that captures the indole attack direction on aminoacrylate. "
    "This blocks the D-Trp screening direction. "
    "Is there someone in the lab who has worked on stereoselectivity-related collective variables?"
)
para(
    "3. GenSLM-230 and NdTrpB structures: I need the sequences and/or homology models "
    "to begin Phase 2 comparisons. Are these available, or should I build models myself?"
)

# ============================================================
# SECTION 5: KEY SIMULATION PARAMETERS
# ============================================================
heading("Section 5: Key Simulation Parameters (from JACS 2019 SI)", level=1)
para(
    "All parameters below were extracted from the Supporting Information of "
    "Maria-Solano et al., JACS 2019 (DOI: 10.1021/jacs.9b03646). "
    "I verified each value against the SI PDF on 2026-03-28."
)
table(
    ["Parameter", "Value", "Source"],
    [
        ["Force field", "AMBER ff14SB (protein-specific parameter set)", "SI Methods"],
        ["Water model", "TIP3P", "SI Methods"],
        ["Box type / buffer", "Cubic, 10 A buffer", "SI Methods"],
        ["Temperature", "350 K", "SI Methods (P. furiosus thermophile)"],
        ["Conventional MD length", "500 ns (NVT production)", "SI Methods"],
        ["Conventional MD engine", "AMBER16 (original); AMBER 24p3 (replication)", "SI Methods / our setup"],
        ["MetaDynamics engine", "GROMACS 5.1.2 + PLUMED 2 (original); GROMACS 2026.0 + PLUMED 2.9 (replication)", "SI Methods / our setup"],
        ["MetaDynamics scheme", "Well-tempered", "SI Methods"],
        ["Gaussian height", "0.15 kcal/mol (0.628 kJ/mol)", "SI Methods"],
        ["Gaussian width", "Adaptive (local FES properties)", "SI Methods"],
        ["Deposition pace", "Every 2 ps (1000 steps at 2 fs timestep)", "SI Methods"],
        ["Bias factor (controls how fast bias decays)", "10", "SI Methods"],
        ["Number of walkers", "10", "SI Table S2"],
        ["Time per walker", "50 to 100 ns", "SI Table S2"],
        ["Path CV atoms (CV = collective variable)", "C-alpha of residues 97-184 + 282-305", "SI Methods"],
        ["Path CV frames", "15 (linear interpolation from open to closed PDB)", "SI Methods"],
        ["Electrostatics", "PME (Particle Mesh Ewald), 8 A cutoff", "SI Methods"],
        ["Bond constraints", "SHAKE (constrains bonds to H); SETTLE (constrains water geometry)", "SI Methods"],
        ["PLP parameterization", "GAFF + RESP at HF/6-31G(d) via antechamber", "SI Methods"],
    ],
)

# ============================================================
# SECTION 6: NEXT WEEK PLAN
# ============================================================
heading("Section 6: Next Week Plan", level=1)

heading("Papers to read", level=2)
para(
    "1. Maria-Solano et al., JACS 2019: full close read with annotation. "
    "I need to understand Figure 2 (FEL comparison across three TrpB systems) "
    "and the productive vs. unproductive closure distinction before evaluating my own results."
)
para(
    "2. Lambert et al., Nature Communications 2026: full read. "
    "Focus on how GenSLM-230 was selected and what experimental assays were used to validate it."
)

heading("Tasks", level=2)
para(
    "1. PLP cofactor parameterization for the Ain intermediate "
    "(the internal aldimine form of the PLP-enzyme complex). "
    "Extract PLP coordinates from PDB 1WDW and compute RESP charges using Gaussian16. "
    "Then generate force field parameter files via antechamber."
)
para(
    "2. Generate 15-frame open-to-closed reference path: "
    "linear interpolation of C-alpha atoms (residues 97-184 + 282-305) "
    "between PDB 1WDW (open) and PDB 3CEP (closed)."
)
para(
    "3. Test GROMACS GPU acceleration: verify that the conda-forge build "
    "supports GPU offloading (important for the 10-walker production runs)."
)
para(
    "4. Begin PfTrpS(Ain) system preparation if PLP parameters are ready: "
    "solvation, energy minimization, 7-step heating to 350 K, "
    "and NPT equilibration (constant pressure and temperature)."
)

# ============================================================
# SECTION 7: REFERENCES
# ============================================================
heading("Section 7: References", level=1)

refs = [
    (
        "1. Maria-Solano, M. A.; Iglesias-Fernandez, J.; Osuna, S. "
        "\"Deciphering the Allosterically Driven Conformational Ensemble in Tryptophan Synthase Evolution.\" "
        "J. Am. Chem. Soc. 2019, 141, 13049-13056. DOI: 10.1021/jacs.9b03646.",
        "Primary source for MetaDynamics protocol, path CV design, and all simulation parameters. "
        "Benchmark target for Phase 1 replication."
    ),
    (
        "2. Maria-Solano, M. A.; Kinateder, T.; Iglesias-Fernandez, J.; Sterner, R.; Osuna, S. "
        "\"In Silico Identification and Experimental Validation of Distal Activity-Enhancing Mutations in TrpB.\" "
        "ACS Catal. 2021, 11, 13733-13743. DOI: 10.1021/acscatal.1c03950.",
        "Extended the JACS 2019 method with experimental validation of SPM-predicted mutations. "
        "Confirms that the MetaDynamics + SPM pipeline has predictive power."
    ),
    (
        "3. Lambert, T.; Tavakoli, A.; Dharuman, G. et al. "
        "\"Sequence-based generative AI design of versatile tryptophan synthases.\" "
        "Nat. Commun. 2026, 17, 1680. DOI: 10.1038/s41467-026-68384-6.",
        "Source of GenSLM-designed TrpB variants. GenSLM-230 is the primary target for Phase 2 comparison."
    ),
    (
        "4. Kinateder, T.; Sterner, R.; Osuna, S. et al. "
        "\"Identification and Characterization of a Naturally Occurring Standalone Beta Subunit of Tryptophan Synthase.\" "
        "Protein Sci. 2025, 34(4), e70103. DOI: 10.1002/pro.70103.",
        "Provides a natural standalone TrpB benchmark (plTrpB). "
        "COMM domain barrier of 5-6 kcal/mol sets a comparative scale for my results."
    ),
]
for ref_text, note in refs:
    para(ref_text)
    para(f"  Role in this project: {note}", italic=True, size=10)

# ============================================================
# SAVE
# ============================================================
outpath = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "WeeklyReport_Week2_2026-03-28.docx")
doc.save(outpath)
print(f"Saved: {outpath}")
print(f"Size: {os.path.getsize(outpath)} bytes")
