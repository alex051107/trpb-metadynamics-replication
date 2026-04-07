"""Generate Week 4 weekly report as .docx"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

doc = Document()

# -- Global style --
style = doc.styles["Normal"]
font = style.font
font.name = "Arial"
font.size = Pt(11)
font.color.rgb = RGBColor(0, 0, 0)
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.2

# Margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


# ---- Helper functions ----

def set_cell_shading(cell, color):
    """Set background shading for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set borders on a table cell. kwargs: top, bottom, left, right with value (color, size, style)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, attrs in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), attrs.get('val', 'single'))
        element.set(qn('w:sz'), str(attrs.get('sz', 4)))
        element.set(qn('w:color'), attrs.get('color', 'D0D0D0'))
        element.set(qn('w:space'), '0')
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_table_borders_light(table, color='D0D0D0'):
    """Set light gray borders on entire table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:color'), color)
        element.set(qn('w:space'), '0')
        borders.append(element)
    # Remove existing borders if any
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def add_paragraph_bottom_border(paragraph, color='CCCCCC', sz='4'):
    """Add a thin bottom border to a paragraph (underline effect)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:color'), color)
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_paragraph_left_border(paragraph, color='2E75B6', sz='12'):
    """Add a left border to a paragraph (accent bar effect)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)
    left.set(qn('w:color'), color)
    left.set(qn('w:space'), '6')
    pBdr.append(left)
    pPr.append(pBdr)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
        if level == 1:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)  # Dark blue
    h.paragraph_format.space_before = Pt(18) if level == 1 else Pt(12)
    if level == 1:
        add_paragraph_bottom_border(h, color='CCCCCC', sz='4')
    return h


def add_header_table(rows_data):
    """Create a metadata-style header table with light gray field column."""
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Remove default style, apply light borders
    table.style = "Table Grid"
    set_table_borders_light(table, color='D0D0D0')

    for r, (field, value) in enumerate(rows_data):
        cell_field = table.rows[r].cells[0]
        cell_value = table.rows[r].cells[1]

        # Field name cell: light gray background, bold
        set_cell_shading(cell_field, 'F2F2F2')
        cell_field.text = field
        for p in cell_field.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(10)

        # Value cell
        cell_value.text = value
        for p in cell_value.paragraphs:
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(10)

        # Light borders on both cells
        border_args = {
            'top': {'color': 'D0D0D0', 'sz': 4},
            'bottom': {'color': 'D0D0D0', 'sz': 4},
            'left': {'color': 'D0D0D0', 'sz': 4},
            'right': {'color': 'D0D0D0', 'sz': 4},
        }
        set_cell_border(cell_field, **border_args)
        set_cell_border(cell_value, **border_args)

    return table


STATUS_COLORS = {
    "Done": RGBColor(0x00, 0x66, 0x00),       # Green
    "Running": RGBColor(0x00, 0x66, 0xCC),     # Blue
    "In Process": RGBColor(0xCC, 0x66, 0x00),  # Orange
}


def add_work_table(headers, rows):
    """Create the work-done table with styled header and alternating rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders_light(table, color='B0B0B0')

    # Header row: dark blue background, white text
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, '2E75B6')
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows with alternating shading
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)

            # Alternating row shading
            if r % 2 == 0:
                set_cell_shading(cell, 'D6E4F0')  # Light blue

            # Style text
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)

            # Color the Status column
            if c == len(headers) - 1 and headers[-1] == "Status":
                status_color = STATUS_COLORS.get(str(val).strip())
                if status_color:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.color.rgb = status_color
                            run.bold = True

    return table


def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.bold = bold
    return p


def add_footer(text):
    """Add footer to all sections."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add page number field
        run = p.add_run(text + " | Page ")
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        # Insert PAGE field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run2 = p.add_run()
        run2._r.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run3 = p.add_run()
        run3._r.append(instrText)
        run3.font.name = "Arial"
        run3.font.size = Pt(9)
        run3.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run4 = p.add_run()
        run4._r.append(fldChar2)


# ============================================================
# TITLE
# ============================================================
title = doc.add_paragraph()
run = title.add_run("Weekly Progress Report")
run.font.size = Pt(22)
run.font.name = "Arial"
run.bold = True
run.font.color.rgb = RGBColor(0, 0, 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.space_after = Pt(4)

subtitle = doc.add_paragraph()
run = subtitle.add_run("TrpB MetaDynamics Replication Project")
run.font.size = Pt(13)
run.font.name = "Arial"
run.font.color.rgb = RGBColor(100, 100, 100)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.space_after = Pt(12)

# ============================================================
# HEADER
# ============================================================
add_header_table([
    ("Date", "April 4, 2026"),
    ("Week", "Week 4: Classical MD Complete + MetaDynamics Pipeline"),
])
doc.add_paragraph()

# ============================================================
# SUMMARY
# ============================================================
add_heading("Summary", level=1)
add_para(
    "The 500 ns classical MD run for PfTrpB(Ain) finished this week. "
    "I then switched from AMBER to GROMACS+PLUMED for the MetaDynamics phase. "
    "ParmEd handled the format conversion without issues."
)
add_para(
    "The majority of my time this week was spent resolving three compatibility issues between GROMACS 2026 and PLUMED 2.9. "
    "Each required extended debugging because the error messages were misleading. "
    "All three are now resolved. "
    "A single-walker MetaDynamics job is running on Longleaf as of today (s=7.8, z=0.5 nm after 55 min). "
    "These early CV values are physically reasonable."
)

# ============================================================
# WORK DONE
# ============================================================
add_heading("Work Done", level=1)
add_work_table(
    ["#", "Task", "Category", "Status"],
    [
        ["1", "500 ns NVT production MD (PfTrpB-Ain, 39268 atoms, Job 40806029)", "Simulation", "Done"],
        ["2", "Group meeting with Dr. Zhang and Amin (April 2)", "Meeting", "Done"],
        ["3", "AMBER-to-GROMACS format conversion via ParmEd", "Setup", "Done"],
        ["4", "Debug FP-018: LAMBDA unit conversion (\u00C5 to nm)", "Debugging", "Done"],
        ["5", "Debug FP-019: GROMACS 2026 PLUMED line-continuation syntax", "Debugging", "Done"],
        ["6", "Debug FP-020: conda PLUMED missing modules; build from source", "Debugging", "Done"],
        ["7", "Submit single-walker MetaDynamics (Job 41500355)", "Simulation", "Running"],
        ["8", "Write TrpB Replication Tutorial (EN + CN)", "Documentation", "In Process"],
    ],
)
doc.add_paragraph()

add_para(
    "Job 40806029 finished in 71.55 hours at 167.72 ns/day on a single GPU. "
    "That gives us the equilibrated PfTrpB(Ain) structure we need as the MetaDynamics starting point."
)
add_para(
    "The AMBER-to-GROMACS format conversion completed without issues. "
    "ParmEd read the topology and restart files and wrote out GROMACS-compatible inputs directly."
)
add_para(
    "For the path collective variable, I interpolated 15 reference structures between the open conformation (1WDW) and the closed conformation (3CEP). "
    "These use 112 C\u03B1 atoms from the COMM domain (residues 97-184) to define the reaction coordinate."
)
add_para(
    "Debugging the three GROMACS-PLUMED compatibility issues occupied Thursday and Friday. "
    "Details are in the Problems section below."
)
add_para(
    "The MetaDynamics simulation is now running with parameters matching the SI protocol. "
    "Adaptive Gaussian widths (ADAPTIVE=GEOM) are enabled. "
    "Early CV values are consistent with expectations."
)
add_para(
    "On April 2 I met with Dr. Zhang and Amin. "
    "Dr. Zhang confirmed the Gaussian input setup for the PLP charge calculation and shared the lab's existing AMBER MD pipeline."
)

# ============================================================
# PROBLEMS ENCOUNTERED
# ============================================================
add_heading("Problems Encountered", level=1)

h = add_heading("Problem 1: Unit mismatch between simulation engines", level=2)
# Add FP label after the heading text
run = h.add_run("  FP-018")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = "Arial"
run.bold = False

add_para(
    "z(R) returned a value of -78, when the expected value was approximately 0.5. "
    "This indicated a problem with the LAMBDA parameter."
)
add_para(
    "AMBER uses angstroms and GROMACS uses nanometers. "
    "I had calculated LAMBDA=0.034 on the AMBER side (in \u00C5\u207B\u00B2) and transferred it directly into the PLUMED input without unit conversion. "
    "This made the exponential kernel 100x too flat, so the path function could not distinguish between reference frames."
)
add_para(
    "This was resolved by multiplying LAMBDA by 100 when converting from \u00C5\u207B\u00B2 to nm\u207B\u00B2 (0.034 becomes 3.39). "
    "After correction, z(R) returned to ~0.5 nm as expected. "
    "I now maintain a unit conversion checklist for any parameter crossing the AMBER-to-GROMACS boundary."
)

h = add_heading("Problem 2: PLUMED input parsing through GROMACS interface", level=2)
run = h.add_run("  FP-019")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = "Arial"
run.bold = False

add_para(
    "PLUMED .dat files normally support backslash line continuation. "
    "Standalone plumed driver handles this correctly. "
    "However, GROMACS 2026 pre-processes the input when loading PLUMED as a plugin and silently drops everything after the backslash."
)
add_para(
    "I encountered a misleading error (\"SIGMA is compulsory\") because SIGMA was on a continuation line that GROMACS truncated. "
    "Initially this suggested ADAPTIVE=GEOM was unsupported, but it functions correctly once all parameters are placed on a single line. "
    "The root cause of why GROMACS handles the backslash differently from standalone PLUMED remains unclear."
)

h = add_heading("Problem 3: Incomplete PLUMED shared library from conda", level=2)
run = h.add_run("  FP-020")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = "Arial"
run.bold = False

add_para(
    "I first attempted to use the conda-forge PLUMED 2.9.2 package. "
    "The command-line tool (plumed driver) worked correctly, but gmx mdrun -plumed failed silently. "
    "After investigation, I determined the issue was in the shared library (libplumedKernel.so) that GROMACS loads at runtime."
)
add_para(
    "The conda build compiles the CLI with all features, but the shared library is missing key modules (PATHMSD and METAD among them). "
    "I compiled PLUMED 2.9.2 from source on Longleaf and set PLUMED_KERNEL to point to the new library, which resolved the issue."
)
add_para(
    "A separate issue arose with PATHMSD and atom indexing. "
    "PATHMSD reads a multi-model PDB with all 15 reference frames and matches atoms by serial number. "
    "Our 112 C\u03B1 atoms have non-sequential serials (1614, 1621, 1643, etc.) scattered across a 39,268-atom system. "
    "Through gmx mdrun, PATHMSD reported zero atoms per frame, though the same file worked correctly in standalone mode."
)
add_para(
    "I switched to FUNCPATHMSD, which uses 15 individual RMSD calculations with single-frame PDB files. "
    "Non-sequential atom serials are handled correctly with this approach. "
    "The path function combines the 15 RMSD values into s(R) and z(R) using the same formula as PATHMSD."
)

# ============================================================
# OPEN QUESTIONS
# ============================================================
add_heading("Open Questions", level=1)

p1 = add_para(
    "1. FUNCPATHMSD vs PATHMSD. "
    "I switched to FUNCPATHMSD because of the atom-indexing issue described in Problem 3. "
    "This was a practical workaround; the underlying math is identical, as s(R) and z(R) are computed the same way."
)
add_paragraph_left_border(p1)

p2 = add_para(
    "I plan to verify whether the separate RMSD alignment steps in FUNCPATHMSD introduce any subtle numerical differences. "
    "I will run a short test using both methods through plumed driver and compare the outputs directly."
)
add_paragraph_left_border(p2)

p3 = add_para(
    "2. MSD discrepancy with the published value. "
    "The SI reports MSD=80 \u00C5\u00B2 between successive path frames (\u03BB=0.029 \u00C5\u207B\u00B2). "
    "I obtain MSD=67.8 \u00C5\u00B2 (\u03BB=0.034 \u00C5\u207B\u00B2), which is 17% higher. "
    "One possible explanation is that the SI does not specify how they aligned 1WDW and 3CEP before interpolation (the two structures come from different species). "
    "Whole-chain alignment instead of COMM-domain-only alignment could bring our MSD closer to 80. "
    "I plan to test this hypothesis."
)
add_paragraph_left_border(p3)

# ============================================================
# NEXT WEEK
# ============================================================
add_heading("Next Week", level=1)
add_para(
    "Job 41500355 should finish in approximately 17 hours. "
    "Once complete, I will run plumed sum_hills to reconstruct the free energy surface. "
    "The priority is to determine whether the O and C basins appear at the expected s(R) positions. "
    "If so, I will proceed to set up the 10-walker production run."
)
add_para(
    "I also plan to review Dr. Zhang's md_setup pipeline, which may simplify future system preparation compared to following the JACS 2019 SI protocol manually."
)
add_para(
    "The Replication Tutorial (EN and CN) remains on hold pending MetaDynamics results."
)

# ============================================================
# FOOTER
# ============================================================
add_footer("Zhenpeng Liu | UNC Chapel Hill")

# Save
outpath = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "..", "WeeklyReport_Week4_2026-04-04.docx")
doc.save(outpath)
print(f"Saved: {outpath}")
print(f"Size: {os.path.getsize(outpath)} bytes")
